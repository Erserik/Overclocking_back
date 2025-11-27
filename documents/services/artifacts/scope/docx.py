import io
from typing import Any, Dict, List

from docx import Document as DocxDocument

DEFAULT_TEXT = "Требует уточнения на основании исходных данных"


def _add_bullets(doc: DocxDocument, items: List[str]) -> None:
    if not items:
        doc.add_paragraph(DEFAULT_TEXT, style="List Bullet")
        return
    for x in items:
        doc.add_paragraph(str(x), style="List Bullet")


def build_docx(structured: Dict[str, Any]) -> bytes:
    """
    Собирает DOCX для Scope на основе structured_data.
    Возвращает байты файла.
    """
    d = DocxDocument()

    summary = (structured.get("summary") or "").strip() or DEFAULT_TEXT

    d.add_heading("Scope", level=1)

    d.add_heading("Summary", level=3)
    d.add_paragraph(summary)

    d.add_heading("In vision", level=3)
    _add_bullets(d, structured.get("in_scope") or [])

    d.add_heading("Out of vision", level=3)
    _add_bullets(d, structured.get("out_of_scope") or [])

    d.add_heading("Business processes in vision", level=3)
    _add_bullets(d, structured.get("business_processes_in_scope") or [])

    d.add_heading("Systems in vision", level=3)
    _add_bullets(d, structured.get("systems_in_scope") or [])

    d.add_heading("Assumptions", level=3)
    _add_bullets(d, structured.get("assumptions") or [])

    d.add_heading("Constraints", level=3)
    _add_bullets(d, structured.get("constraints") or [])

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
