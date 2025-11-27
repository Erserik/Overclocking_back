import io
from typing import Any, Dict, List

from docx import Document as DocxDocument

DEFAULT_TEXT = "Требует уточнения на основании исходных данных"


def _add_bullets(doc: DocxDocument, items: List[str]) -> None:
    if not items:
        doc.add_paragraph(DEFAULT_TEXT, style="List Bullet")
        return
    for x in items:
        doc.add_paragraph(str(x), style="List Bullet")


def build_docx(structured: Dict[str, Any]) -> bytes:
    """
    Собирает DOCX для Vision на основе structured_data.
    Возвращает байты файла.
    """
    d = DocxDocument()

    title = (structured.get("title") or "").strip() or "Vision"
    problem = (structured.get("problem_statement") or "").strip() or DEFAULT_TEXT

    d.add_heading("Vision", level=1)
    d.add_heading(title, level=2)

    d.add_heading("Problem statement", level=3)
    d.add_paragraph(problem)

    d.add_heading("Business goals", level=3)
    _add_bullets(d, structured.get("business_goals") or [])

    d.add_heading("Target users", level=3)
    _add_bullets(d, structured.get("target_users") or [])

    d.add_heading("Expected outcomes", level=3)
    _add_bullets(d, structured.get("expected_outcomes") or [])

    d.add_heading("Success criteria", level=3)
    _add_bullets(d, structured.get("success_criteria") or [])

    d.add_heading("Risks and limitations", level=3)
    _add_bullets(d, structured.get("risks_and_limitations") or [])

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
